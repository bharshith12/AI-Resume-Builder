import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any

def generate_docx_resume(resume_data: Dict[str, Any]) -> io.BytesIO:
    doc = Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    personal = resume_data.get("personalInfo", {})
    summary = resume_data.get("summary", "")
    experience = resume_data.get("experience", [])
    education = resume_data.get("education", [])
    skills = resume_data.get("skills", {})
    projects = resume_data.get("projects", [])

    # Header Name
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run(personal.get("fullName", "Your Name").upper())
    run_name.bold = True
    run_name.font.size = Pt(20)
    run_name.font.color.rgb = RGBColor(30, 41, 59)

    # Contact Line
    contact_parts = []
    if personal.get("email"): contact_parts.append(personal.get("email"))
    if personal.get("phone"): contact_parts.append(personal.get("phone"))
    if personal.get("address"): contact_parts.append(personal.get("address"))
    if personal.get("linkedin"): contact_parts.append(personal.get("linkedin"))

    p_contact = doc.add_paragraph(" | ".join(contact_parts))
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.style.font.size = Pt(9.5)

    def add_section_header(title_text):
        p_hdr = doc.add_paragraph()
        p_hdr.paragraph_format.space_before = Pt(12)
        p_hdr.paragraph_format.space_after = Pt(4)
        run = p_hdr.add_run(title_text.upper())
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(15, 23, 42)

    # Summary
    if summary:
        add_section_header("Professional Summary")
        p_sum = doc.add_paragraph(summary)
        p_sum.style.font.size = Pt(10)

    # Experience
    if experience:
        add_section_header("Experience")
        for exp in experience:
            p_exp = doc.add_paragraph()
            r_pos = p_exp.add_run(f"{exp.get('position', 'Role')} - {exp.get('company', 'Company')}")
            r_pos.bold = True
            r_pos.font.size = Pt(10.5)

            dates = f" ({exp.get('startDate', '')} - {exp.get('endDate', 'Present')})"
            r_date = p_exp.add_run(dates)
            r_date.italic = True
            r_date.font.size = Pt(9.5)

            bullets = exp.get("bullets", [])
            if not bullets and exp.get("description"):
                bullets = [b.strip() for b in exp.get("description").split(".") if b.strip()]

            for bullet in bullets:
                p_b = doc.add_paragraph(bullet, style='List Bullet')
                p_b.style.font.size = Pt(9.5)
                p_b.paragraph_format.space_after = Pt(2)

    # Projects
    if projects:
        add_section_header("Projects")
        for proj in projects:
            p_proj = doc.add_paragraph()
            r_title = p_proj.add_run(f"{proj.get('title', 'Project Title')}")
            r_title.bold = True
            r_title.font.size = Pt(10.5)

            if proj.get("technologies"):
                tech_str = f" | {', '.join(proj.get('technologies')) if isinstance(proj.get('technologies'), list) else proj.get('technologies')}"
                r_tech = p_proj.add_run(tech_str)
                r_tech.italic = True

            if proj.get("description"):
                p_desc = doc.add_paragraph(proj.get("description"), style='List Bullet')
                p_desc.style.font.size = Pt(9.5)

    # Education
    if education:
        add_section_header("Education")
        for edu in education:
            p_edu = doc.add_paragraph()
            r_deg = p_edu.add_run(f"{edu.get('degree', 'Degree')}, {edu.get('field', '')}")
            r_deg.bold = True
            r_school = p_edu.add_run(f" - {edu.get('institution', 'University')}")
            r_school.font.size = Pt(10)

    # Skills
    if skills:
        add_section_header("Skills")
        p_sk = doc.add_paragraph()
        if isinstance(skills, dict):
            tech_s = ", ".join(skills.get("technical", []))
            soft_s = ", ".join(skills.get("soft", []))
            if tech_s:
                p_sk.add_run("Technical: ").bold = True
                p_sk.add_run(tech_s + "\n")
            if soft_s:
                p_sk.add_run("Soft Skills: ").bold = True
                p_sk.add_run(soft_s)
        elif isinstance(skills, list):
            p_sk.add_run(", ".join([str(s) for s in skills]))

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
